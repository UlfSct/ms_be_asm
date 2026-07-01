#!/usr/bin/env python3
"""
Скрипт для сбора исходного кода Django проекта в DOCX файл.
Исключает служебные директории (venv, migrations, static и т.д.)
"""

import sys
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import argparse

# Директории, которые нужно исключить
EXCLUDED_DIRS = {
    # Общие
    'KrasheMit',
    'node_modules',
    'dist',
    '.git',
    '.svn',
    '.hg',
    '.idea',
    '.vscode',
    '.vs',
    'coverage',
    '.nyc_output',
    'build',
    '.next',
    '.nuxt',
    '.output',
    'target',
    'vendor',
    'bower_components',
    '.cache',
    '.parcel-cache',
    '.turbo',

    # Python/Django специфичные
    'venv',
    '.venv',
    'env',
    '.env',
    'virtualenv',
    '__pycache__',
    '.pytest_cache',
    '.mypy_cache',
    '.tox',
    'migrations',
    'static',
    'staticfiles',
    'media',
    'mediafiles',
    'collected_static',
}

EXCLUDED_FILES = {
    # Общие
    'package-lock.json',
    'yarn.lock',
    'pnpm-lock.yaml',
    'bun.lockb',
    '.DS_Store',
    'Thumbs.db',
    '.env',
    '.env.local',
    '.env.development',
    '.env.production',
    '.env.test',

    # Python/Django специфичные
    'db.sqlite3',
    'db.sqlite3-journal',
    '*.pyc',
    '*.pyo',
    'celerybeat-schedule',
    'celerybeat.pid',
}

# Расширения файлов для включения
INCLUDED_EXTENSIONS = {
    # Python
    '.py',

    # JavaScript/TypeScript
    '.js', '.ts', '.vue', '.jsx', '.tsx',

    # Стили
    '.css', '.scss', '.sass', '.less',

    # Шаблоны
    '.html', '.htm', '.jinja', '.jinja2', '.djhtml',

    # Конфигурация
    '.json', '.yaml', '.yml', '.toml', '.cfg', '.ini', '.conf',

    # Документация
    '.md', '.txt', '.rst', '.adoc',

    # Другие
    '.xml', '.svg',
    '.graphql', '.gql',
    '.env.example', '.env.sample',
    '.prisma',
    '.eslintrc', '.prettierrc', '.babelrc',
    '.editorconfig', '.gitignore', '.npmignore', '.dockerignore',

    # Requirements
    '.txt',  # для requirements файлов
}

# Важные конфигурационные файлы
IMPORTANT_CONFIGS = {
    # Django основные
    'settings.py',
    'urls.py',
    'wsgi.py',
    'asgi.py',
    'manage.py',

    # Python пакеты
    'requirements.txt',
    'requirements-dev.txt',
    'requirements-prod.txt',
    'requirements-test.txt',
    'Pipfile',
    'Pipfile.lock',
    'pyproject.toml',
    'setup.py',
    'setup.cfg',
    'tox.ini',

    # Тестирование и качество кода
    'pytest.ini',
    'pytest.cfg',
    '.coveragerc',
    '.flake8',
    '.pylintrc',
    'mypy.ini',
    '.isort.cfg',
    'pre-commit-config.yaml',

    # Деплой
    'Dockerfile',
    'docker-compose.yml',
    'docker-compose.yaml',
    '.dockerignore',
    'Procfile',
    'runtime.txt',
    'nginx.conf',
    'uwsgi.ini',
    'gunicorn.conf.py',

    # Документация
    'README.md',
    'CHANGELOG.md',
    'LICENSE',
    'CONTRIBUTING.md',
    'AUTHORS',

    # Git и редактор
    '.gitignore',
    '.editorconfig',
    '.env.example',
    '.env.sample',

    # Makefile
    'Makefile',
    'makefile',
}


def is_config_file(filename):
    """Проверяет, является ли файл важным конфигом"""
    return filename in IMPORTANT_CONFIGS


def should_include_file(filepath):
    """Определяет, нужно ли включать файл в документ"""
    filename = filepath.name

    # Проверяем исключённые файлы
    if filename in EXCLUDED_FILES:
        return False

    # Проверяем важные конфиги
    if is_config_file(filename):
        return True

    # Проверяем расширение
    suffix = filepath.suffix
    if suffix in INCLUDED_EXTENSIONS:
        return True

    # Проверяем составные расширения (.env.example и т.д.)
    if filepath.suffixes and ''.join(filepath.suffixes) in INCLUDED_EXTENSIONS:
        return True

    # Проверяем файлы без расширения, которые являются конфигами
    if not suffix and filename in IMPORTANT_CONFIGS:
        return True

    return False


def should_exclude_dir(dirpath):
    """Проверяет, нужно ли исключить директорию"""
    return dirpath.name in EXCLUDED_DIRS or dirpath.name.startswith('.')


def collect_files(root_path):
    """Собирает все файлы проекта, исключая служебные"""
    root = Path(root_path).resolve()

    if not root.exists():
        print(f"Ошибка: путь '{root_path}' не существует")
        sys.exit(1)

    if not root.is_dir():
        print(f"Ошибка: '{root_path}' не является директорией")
        sys.exit(1)

    collected_files = []

    for dirpath, dirnames, filenames in os.walk(root):
        dirpath = Path(dirpath)

        # Исключаем служебные директории из обхода
        dirnames[:] = [
            d for d in dirnames
            if not should_exclude_dir(Path(d))
        ]

        for filename in filenames:
            filepath = dirpath / filename

            if should_include_file(filepath):
                relative_path = filepath.relative_to(root)
                collected_files.append(relative_path)

    return sorted(collected_files, key=lambda p: str(p).lower())


def add_code_block(doc, filepath, content):
    """Добавляет блок кода в документ"""
    # Пустая строка перед названием файла (с минимальным интервалом)
    spacer = doc.add_paragraph()
    pf = spacer.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(1)

    # Название файла жирным и подчёркнутым
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.add_run(str(filepath))
    run.bold = True
    run.underline = True
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    pf = heading.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(10)

    # Пустая строка после названия файла (с минимальным интервалом)
    spacer2 = doc.add_paragraph()
    pf = spacer2.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(1)

    # Код
    try:
        # Убираем пустые строки в конце файла
        content = content.rstrip('\n\r')
        lines = content.split('\n')

        for line in lines:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Код строки (заменяем табы на пробелы для читаемости)
            code_run = para.add_run(line.replace('\t', '    '))
            code_run.font.name = 'Courier New'
            code_run.font.size = Pt(8)

            # Настройка межстрочного интервала
            pf = para.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing = Pt(10)
    except Exception as e:
        error_para = doc.add_paragraph()
        error_run = error_para.add_run(f"[Ошибка чтения файла: {e}]")
        error_run.font.name = 'Courier New'
        error_run.font.size = Pt(8)

    # Пустая строка для разделения между файлами (с минимальным интервалом)
    spacer3 = doc.add_paragraph()
    pf = spacer3.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(0)
    pf.line_spacing = Pt(1)


def read_file_content(root_path, relative_path):
    """Читает содержимое файла"""
    full_path = root_path / relative_path

    # Пропускаем бинарные файлы
    binary_extensions = {'.woff', '.woff2', '.ttf', '.eot', '.otf',
                         '.png', '.jpg', '.jpeg', '.gif', '.ico', '.webp',
                         '.mp3', '.mp4', '.wav', '.ogg', '.pdf',
                         '.sqlite3', '.db', '.sqlite'}

    if full_path.suffix.lower() in binary_extensions:
        return f"[Бинарный файл: {full_path.stat().st_size} байт]"

    try:
        with open(full_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()

        # Ограничиваем размер (чтобы не создавать огромный документ)
        max_size = 100000  # 100KB
        if len(content) > max_size:
            content = content[:max_size] + f"\n\n... [Файл обрезан. Полный размер: {len(content)} символов]"

        return content
    except Exception as e:
        return f"[Ошибка чтения: {str(e)}]"


def create_document(root_path, output_path=None):
    """Создаёт DOCX документ со всем кодом проекта"""
    root = Path(root_path).resolve()

    if output_path is None:
        project_name = root.name or 'django-project'
        # Сохраняем в текущую рабочую директорию (где запущен скрипт)
        output_path = Path.cwd() / f"{project_name}-source-code.docx"
    else:
        output_path = Path(output_path)

    print(f"Сканирование проекта: {root}")
    files = collect_files(root)

    if not files:
        print("Не найдено файлов для включения в документ")
        sys.exit(1)

    print(f"Найдено файлов: {len(files)}")

    # Создаём документ
    doc = Document()

    # Настройка стиля по умолчанию
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Courier New'
    font.size = Pt(8)

    # Настройка страницы
    section = doc.sections[0]
    section.page_width = Cm(29.7)  # A4 landscape
    section.page_height = Cm(21.0)

    # Настройка двух колонок
    sectPr = section._sectPr
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')  # 2 колонки
    cols.set(qn('w:space'), '720')  # Отступ между колонками
    sectPr.append(cols)

    # Настройка полей
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

    # Добавляем файлы
    print("Добавление файлов в документ...")
    for i, filepath in enumerate(files, 1):
        print(f"  [{i}/{len(files)}] {filepath}")

        content = read_file_content(root, filepath)
        add_code_block(doc, filepath, content)

    # Сохраняем документ
    doc.save(str(output_path))
    print(f"\nДокумент сохранён: {output_path}")
    print(f"Размер файла: {output_path.stat().st_size / 1024:.1f} KB")


def main():
    parser = argparse.ArgumentParser(
        description='Сбор исходного кода Django проекта в DOCX файл с двумя колонками'
    )
    parser.add_argument(
        'path',
        type=str,
        help='Абсолютный путь к корню Django проекта'
    )
    parser.add_argument(
        '-o', '--output',
        type=str,
        default=None,
        help='Путь для сохранения выходного DOCX файла (по умолчанию: текущая директория)'
    )

    args = parser.parse_args()

    create_document(args.path, args.output)


if __name__ == '__main__':
    main()
    #  python.exe .\collect_back_code_to_docx.py "C:\Users\anto8\OneDrive\рабочий стол\диплом\be_asm_3d"
